"""Build the integrated Adap manuscript from the archived frozen evidence.

The archived language is already frozen in ``manuscript_content.py``. This
portable builder has no dependency on an external editing source or project
directory and never overwrites the audited manuscript copies.
"""

from __future__ import annotations

import hashlib
import re
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
OUT = HERE.parent / "build"
sys.path.insert(0, str(HERE))

from manuscript_content import (  # noqa: E402
    ABSTRACT,
    AI_DISCLOSURE,
    CONCLUSION,
    FALSIFICATION_ROWS,
    KEYWORDS,
    SECTIONS,
    TITLE,
)
from work.icceic_paper_content import REFERENCES  # noqa: E402
from work.paper_data import EXPECTED_METHODS, validate_contract  # noqa: E402


TEMPLATE = HERE / "template" / "Conference-template-letter-transitional.docx"
PIPELINE_FIGURE = HERE / "assets" / "adap_pipeline.png"

OUTPUTS = {
    "anonymous": OUT / "AdapTabPrompt_ICCEIC2026_Final_Anonymous.docx",
    "author": OUT / "AdapTabPrompt_ICCEIC2026_Final_Author_Copy.docx",
}

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
EP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest().upper()


def clean_text(text: str) -> str:
    """Normalize fragile punctuation while retaining names and mathematical meaning."""
    return (
        text.replace("\u2014", "-")
        .replace("\u2013", "-")
        .replace("\u2011", "-")
        .replace("\u00a0", " ")
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
    )


def set_font(run, size=10.0, bold=None, italic=None, name="Times New Roman", gray=0):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor(gray, gray, gray)


def configure_page(section, columns=1):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    cols = section._sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        section._sectPr.append(cols)
    cols.set(qn("w:num"), str(columns))
    cols.set(qn("w:space"), "360")
    if columns == 2:
        cols.set(qn("w:equalWidth"), "1")
    else:
        cols.attrib.pop(qn("w:equalWidth"), None)


def clear_template(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 0.99
    pf.space_after = Pt(0.5)
    pf.first_line_indent = Inches(0.14)
    pf.widow_control = True


def add_paragraph(doc, text, *, size=10.0, indent=True, italic=False, keep=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 0.99
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.first_line_indent = Inches(0.14) if indent else Inches(0)
    p.paragraph_format.widow_control = True
    p.paragraph_format.keep_together = keep
    set_font(p.add_run(clean_text(text)), size=size, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(2 if level == 1 else 0)
    p.paragraph_format.first_line_indent = Inches(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    set_font(p.add_run(clean_text(text)), size=10, bold=(level == 1), italic=(level == 2))
    return p


def set_cell_shading(cell, fill="E6E6E6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, value=28):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    total_dxa = int(round(sum(widths) * 1440))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width_dxa = int(round(widths[index] * 1440))
            cell.width = Inches(widths[index])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width_dxa))


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "bottom", "insideH"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6" if edge != "insideH" else "3")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "666666")
    for edge in ("left", "right", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def add_table(doc, caption, headers, rows, widths, body_size=7.5):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(1)
    cap.paragraph_format.keep_with_next = True
    cap.paragraph_format.first_line_indent = Inches(0)
    set_font(cap.add_run(clean_text(caption)), size=8, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    header = table.rows[0]
    header._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.line_spacing = 0.92
        set_font(p.add_run(clean_text(str(value))), size=body_size, bold=True)
    for values in rows:
        row = table.add_row()
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.line_spacing = 0.90
            set_font(p.add_run(clean_text(str(value))), size=body_size)
    set_table_geometry(table, widths)
    return table


def one_column(doc):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_page(section, 1)


def two_columns(doc):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_page(section, 2)


def add_title_band(doc, named=False):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    title.paragraph_format.first_line_indent = Inches(0)
    set_font(title.add_run(TITLE), size=24)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(2 if named else 5)
    author.paragraph_format.keep_with_next = True
    author.paragraph_format.first_line_indent = Inches(0)
    set_font(author.add_run("Lin Zhanyi" if named else "Anonymous Authors"), size=11)
    if named:
        affiliation = doc.add_paragraph()
        affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
        affiliation.paragraph_format.space_after = Pt(5)
        affiliation.paragraph_format.keep_with_next = True
        affiliation.paragraph_format.first_line_indent = Inches(0)
        set_font(
            affiliation.add_run(
                "Hong Kong Metropolitan University, lzzzy20041125@outlook.com"
            ),
            size=10,
        )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.19)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 0.98
    set_font(p.add_run("Abstract-"), size=9, bold=True)
    set_font(p.add_run(ABSTRACT), size=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.19)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 0.98
    set_font(p.add_run("Keywords-"), size=9, bold=True, italic=True)
    set_font(p.add_run(KEYWORDS), size=9, italic=True)


def add_pipeline_figure(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_together = True
    p.paragraph_format.first_line_indent = Inches(0)
    p.add_run().add_picture(str(PIPELINE_FIGURE), width=Inches(3.30))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.first_line_indent = Inches(0)
    cap.paragraph_format.space_after = Pt(2)
    cap.paragraph_format.keep_together = True
    set_font(
        cap.add_run(
            "Figure 1. Registered adaptive rescue flow. Candidate features may enter the "
            "fixed semantic cache, but candidate and evaluation labels never enter "
            "prototypes, gating, thresholding, or rescue. A rejected or empty rescue "
            "returns the anchor exactly."
        ),
        size=8,
    )


def grouped_rows(data):
    grouped = {}
    for row in data["method_rows"]:
        key = (row["dataset"], row["shot"], row["seed"])
        grouped.setdefault(key, {})[row["method"]] = row
    return grouped


def episode_rows(data):
    grouped = grouped_rows(data)
    rows = []
    for dataset in ("credit-g", "spambase"):
        for shot in (5, 10, 20):
            for seed in (0, 9, 19):
                g = grouped[(dataset, shot, seed)]
                values = [g[m] for m in EXPECTED_METHODS]
                rows.append(
                    (
                        dataset,
                        shot,
                        seed,
                        f"{values[0]['roc_auc']:.5f}",
                        f"{values[1]['roc_auc']:.5f}",
                        f"{values[2]['roc_auc']:.5f}",
                        f"{values[3]['roc_auc']:.5f}",
                        f"{values[1]['gate_accepted']}/{values[1]['rescued_count']}",
                        f"{values[2]['gate_accepted']}/{values[2]['rescued_count']}",
                        f"{values[3]['gate_accepted']}/{values[3]['rescued_count']}",
                    )
                )
    return rows


def decision_rows(data):
    report = data["report"]
    stats = report["statistics"]
    conjuncts = report["conjuncts"]
    specs = [
        ("Overall Main-Tree mean", ">= +0.0200", f"{stats['overall_main_minus_tree_mean']:+.4f}", "overall_main_delta_min"),
        ("Main-NoSem mean", ">= +0.0100", f"{stats['main_minus_no_semantics_mean']:+.4f}", "main_minus_no_semantics_min"),
        ("Main-Perm mean", ">= +0.0100", f"{stats['main_minus_permuted_mean']:+.4f}", "main_minus_permuted_min"),
        ("Dataset mean/median", "both nonnegative", "both 0.0000", "both_dataset_mean_and_median_nonnegative"),
        ("Positive shot means", ">= 2 of 3", str(stats["positive_shot_mean_count"]), "positive_shot_means_min_count"),
        ("Rescue coverage", "both datasets", "neither (main)", "rescue_required_both_datasets"),
        ("Pseudo-label collapse", "zero", "zero", "zero_collapse"),
        ("Failed method rows", "zero", str(report["failure_row_count"]), "no_failed_method_rows"),
        ("Exact registered rows", "18/72", f"{report['episode_count']}/{report['method_row_count']}", "exact_registered_rows"),
        ("Exact fallback", "all fallback exact", "pass", "all_fallback_and_alpha_zero_exact"),
    ]
    return [
        (condition, threshold, observed, "PASS" if conjuncts[key] else "FAIL")
        for condition, threshold, observed, key in specs
    ]


def add_references(doc):
    doc.add_page_break()
    add_heading(doc, "REFERENCES")
    for index, reference in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.17)
        p.paragraph_format.first_line_indent = Inches(-0.17)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 0.91
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_font(p.add_run(clean_text(reference)), size=8)
        # Keep the final IEEE reference page visually balanced.  The retained
        # template does not balance a short terminal two-column section under
        # Word or WPS, so place references 1-8 and 9-15 explicitly.
        if index == 8:
            p.add_run().add_break(WD_BREAK.COLUMN)
    # A terminal continuous section break makes Word balance the two reference
    # columns on the final page instead of leaving the second column empty.
    one_column(doc)


def remove_privacy_and_revision_parts(path: Path, metadata_author: str):
    with tempfile.TemporaryDirectory() as temp_name:
        temp = Path(temp_name)
        with zipfile.ZipFile(path) as source:
            source.extractall(temp)

        victims = [
            temp / "docProps" / "custom.xml",
            temp / "word" / "comments.xml",
            temp / "word" / "commentsExtended.xml",
            temp / "word" / "people.xml",
        ]
        for victim in victims:
            if victim.exists():
                victim.unlink()

        for xml_path in (temp / "word").rglob("*.xml"):
            tree = ET.parse(xml_path)
            root = tree.getroot()
            changed = False
            for node in root.iter():
                for attr in list(node.attrib):
                    if attr.startswith(f"{{{W_NS}}}rsid"):
                        del node.attrib[attr]
                        changed = True
            if changed:
                tree.write(xml_path, encoding="utf-8", xml_declaration=True)

        for rel_path in [temp / "_rels" / ".rels", temp / "word" / "_rels" / "document.xml.rels"]:
            if not rel_path.exists():
                continue
            tree = ET.parse(rel_path)
            root = tree.getroot()
            for node in list(root):
                target = (node.get("Target") or "").lower()
                rel_type = (node.get("Type") or "").lower()
                if any(token in target or token in rel_type for token in ("comments", "people.xml", "custom-properties", "docprops/custom.xml")):
                    root.remove(node)
            tree.write(rel_path, encoding="utf-8", xml_declaration=True)

        content_types = temp / "[Content_Types].xml"
        if content_types.exists():
            tree = ET.parse(content_types)
            root = tree.getroot()
            for node in list(root):
                value = ((node.get("PartName") or "") + " " + (node.get("ContentType") or "")).lower()
                if any(token in value for token in ("comments", "people+xml", "custom-properties", "/docprops/custom.xml")):
                    root.remove(node)
            tree.write(content_types, encoding="utf-8", xml_declaration=True)

        core = temp / "docProps" / "core.xml"
        if core.exists():
            tree = ET.parse(core)
            root = tree.getroot()
            for tag in (f"{{{DC_NS}}}creator", f"{{{CP_NS}}}lastModifiedBy"):
                node = root.find(tag)
                if node is not None:
                    node.text = metadata_author
            for tag in (f"{{{DC_NS}}}subject", f"{{{DC_NS}}}description"):
                node = root.find(tag)
                if node is not None:
                    node.text = ""
            tree.write(core, encoding="utf-8", xml_declaration=True)

        app = temp / "docProps" / "app.xml"
        if app.exists():
            tree = ET.parse(app)
            root = tree.getroot()
            for tag in ("Company", "Manager"):
                node = root.find(f"{{{EP_NS}}}{tag}")
                if node is not None:
                    node.text = ""
            tree.write(app, encoding="utf-8", xml_declaration=True)

        rebuilt = path.with_suffix(".rebuilt.docx")
        with zipfile.ZipFile(rebuilt, "w", zipfile.ZIP_DEFLATED) as target:
            for item in sorted(temp.rglob("*")):
                if item.is_file():
                    target.write(item, item.relative_to(temp).as_posix())
        rebuilt.replace(path)


def build(mode: str, data):
    named = mode == "author"
    output = OUTPUTS[mode]
    doc = Document(TEMPLATE)
    clear_template(doc)
    configure_styles(doc)
    configure_page(doc.sections[0], 1)
    add_title_band(doc, named=named)
    two_columns(doc)

    for section_title, paragraphs in SECTIONS:
        add_heading(doc, section_title)
        for text in paragraphs:
            add_paragraph(doc, text)

        if section_title.startswith("III."):
            add_table(
                doc,
                "TABLE I. MATCHED INFORMATION ACCESS",
                ("Route", "Semantic content", "Candidate labels", "Endpoint"),
                (
                    ("Tree", "No", "No", "Tree"),
                    ("Main", "Fixed T0-3B", "No", "Tree after accepted rescue or exact fallback"),
                    ("NoSem", "Removed", "No", "Matched control tree"),
                    ("Perm", "Row-permuted", "No", "Matched control tree"),
                ),
                widths=(0.45, 0.70, 0.65, 1.70),
                body_size=7.2,
            )
            add_pipeline_figure(doc)

        if section_title.startswith("IV."):
            one_column(doc)
            add_table(
                doc,
                "TABLE II. COMPLETE FROZEN EPISODE LEDGER (ROC-AUC AND GATE/RESCUE COUNTS)",
                ("Dataset", "Shot", "Seed", "Tree", "Main", "NoSem", "Perm", "M G/R", "N G/R", "P G/R"),
                episode_rows(data),
                widths=(1.00, 0.50, 0.50, 0.80, 0.80, 0.80, 0.80, 0.68, 0.68, 0.68),
                body_size=7.2,
            )
            note = doc.add_paragraph()
            note.paragraph_format.first_line_indent = Inches(0)
            note.paragraph_format.space_after = Pt(2)
            note.paragraph_format.line_spacing = 0.92
            set_font(
                note.add_run(
                    "Tree = TreeSelfTrain-S; Main = ProtoOOF-Tree-v1; NoSem = "
                    "ProtoOOF-NoSemantics; Perm = ProtoOOF-Permuted; G/R = gate "
                    "accepted/rescued rows. All 72 method rows succeeded; all collapse "
                    "indicators are zero."
                ),
                size=7.2,
            )
            two_columns(doc)

        if section_title.startswith("V."):
            add_table(
                doc,
                "TABLE III. REGISTERED CONJUNCTIVE DECISION",
                ("Condition", "Threshold", "Observed", "Result"),
                decision_rows(data),
                widths=(1.15, 0.72, 0.93, 0.70),
                body_size=7.1,
            )

        if section_title.startswith("VIII."):
            add_table(
                doc,
                "TABLE IV. FALSIFIABLE FOLLOW-UP HYPOTHESES",
                ("Hypothesis", "Evidence needed to challenge the boundary"),
                FALSIFICATION_ROWS,
                widths=(1.10, 2.40),
                body_size=7.2,
            )

    add_heading(doc, "IX. CONCLUSION")
    add_paragraph(doc, CONCLUSION)
    add_paragraph(doc, AI_DISCLOSURE, size=8.0, indent=False, keep=True)
    add_references(doc)

    metadata_author = "Lin Zhanyi" if named else "Anonymous"
    doc.core_properties.title = TITLE
    doc.core_properties.author = metadata_author
    doc.core_properties.last_modified_by = metadata_author
    doc.core_properties.subject = ""
    doc.core_properties.comments = ""
    doc.core_properties.keywords = ""
    doc.core_properties.category = ""
    doc.save(output)
    remove_privacy_and_revision_parts(output, metadata_author)
    return output


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    template_before = sha256(TEMPLATE)
    data = validate_contract()
    built = {mode: build(mode, data) for mode in OUTPUTS}
    if sha256(TEMPLATE) != template_before:
        raise SystemExit("official template changed unexpectedly")
    for mode, path in built.items():
        print(f"{mode}: {path} {sha256(path)}")


if __name__ == "__main__":
    main()
