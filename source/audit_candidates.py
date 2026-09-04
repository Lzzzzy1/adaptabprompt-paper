from __future__ import annotations

import hashlib
import json
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from pypdf import PdfReader


SOURCE = Path(__file__).resolve().parent
ROOT = SOURCE.parent / "manuscript"
sys.path.insert(0, str(SOURCE))

from work.paper_data import EXPECTED_METHODS, validate_contract  # noqa: E402


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CP = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC = "http://purl.org/dc/elements/1.1/"
EP = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
NS = {"w": W, "cp": CP, "dc": DC, "ep": EP}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest().upper()


def clean_cell(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def expected_episode_rows(data: dict) -> list[list[str]]:
    grouped: dict[tuple[str, int, int], dict[str, dict]] = {}
    for row in data["method_rows"]:
        grouped.setdefault((row["dataset"], row["shot"], row["seed"]), {})[
            row["method"]
        ] = row
    rows = []
    for dataset in ("credit-g", "spambase"):
        for shot in (5, 10, 20):
            for seed in (0, 9, 19):
                values = [grouped[(dataset, shot, seed)][method] for method in EXPECTED_METHODS]
                rows.append(
                    [
                        dataset,
                        str(shot),
                        str(seed),
                        f"{values[0]['roc_auc']:.5f}",
                        f"{values[1]['roc_auc']:.5f}",
                        f"{values[2]['roc_auc']:.5f}",
                        f"{values[3]['roc_auc']:.5f}",
                        f"{values[1]['gate_accepted']}/{values[1]['rescued_count']}",
                        f"{values[2]['gate_accepted']}/{values[2]['rescued_count']}",
                        f"{values[3]['gate_accepted']}/{values[3]['rescued_count']}",
                    ]
                )
    return rows


def expected_decision_rows(data: dict) -> list[list[str]]:
    report = data["report"]
    stats = report["statistics"]
    conjuncts = report["conjuncts"]
    specs = [
        ("Overall Main-Tree mean", ">= +0.0200", f"{stats['overall_main_minus_tree_mean']:+.4f}", "overall_main_delta_min"),
        ("Main-NoSem mean", ">= +0.0100", f"{stats['main_minus_no_semantics_mean']:+.4f}", "main_minus_no_semantics_min"),
        ("Main-Perm mean", ">= +0.0100", f"{stats['main_minus_permuted_mean']:+.4f}", "main_minus_permuted_min"),
        ("Dataset mean/median", "both nonnegative", "both 0.0000", "both_dataset_mean_and_median_nonnegative"),
        ("Positive shot means", ">= 2 of 3", str(stats["positive_shot_mean_count"]), "positive_shot_means_min_count"),
        ("Rescue coverage", "both datasets", "neither (main)", "rescue_required_both_datasets"),
        ("Pseudo-label collapse", "zero", "zero", "zero_collapse"),
        ("Failed method rows", "zero", str(report["failure_row_count"]), "no_failed_method_rows"),
        ("Exact registered rows", "18/72", f"{report['episode_count']}/{report['method_row_count']}", "exact_registered_rows"),
        ("Exact fallback", "all fallback exact", "pass", "all_fallback_and_alpha_zero_exact"),
    ]
    return [
        [condition, threshold, observed, "PASS" if conjuncts[key] else "FAIL"]
        for condition, threshold, observed, key in specs
    ]


def expand_citations(text: str) -> set[int]:
    found: set[int] = set()
    for match in re.finditer(r"\[([^\]]+)\]", text):
        content = match.group(1)
        if not re.fullmatch(r"[\d,\-\s]+", content):
            continue
        for part in content.split(","):
            part = part.strip()
            if "-" in part:
                start, end = (int(value) for value in part.split("-", 1))
                found.update(range(start, end + 1))
            elif part:
                found.add(int(part))
    return found


def xml_text(root: ET.Element) -> str:
    return "".join(node.text or "" for node in root.iter(f"{{{W}}}t"))


def docx_audit(path: Path, data: dict) -> dict:
    result: dict = {"path": str(path), "sha256": sha256(path), "size": path.stat().st_size}
    with zipfile.ZipFile(path) as archive:
        result["zip_test"] = archive.testzip()
        names = set(archive.namelist())
        result["comment_parts"] = sorted(name for name in names if "comment" in name.lower() or "people.xml" in name.lower())
        result["custom_props"] = "docProps/custom.xml" in names

        story_names = [
            name
            for name in names
            if re.fullmatch(r"word/(document|header\d*|footer\d*|footnotes|endnotes)\.xml", name)
        ]
        tracked = comments = highlights = rsids = 0
        colors: list[dict] = []
        story_text = ""
        for name in story_names:
            root = ET.fromstring(archive.read(name))
            story_text += "\n" + xml_text(root)
            tracked += sum(len(root.findall(f".//w:{tag}", NS)) for tag in ("ins", "del", "moveFrom", "moveTo"))
            comments += sum(
                len(root.findall(f".//w:{tag}", NS))
                for tag in ("commentRangeStart", "commentRangeEnd", "commentReference")
            )
            highlights += len(root.findall(".//w:highlight", NS))
            rsids += sum(
                1
                for node in root.iter()
                for attr in node.attrib
                if attr.startswith(f"{{{W}}}rsid")
            )
            for node in root.findall(".//w:color", NS):
                colors.append(
                    {
                        "part": name,
                        "val": node.get(f"{{{W}}}val"),
                        "theme": node.get(f"{{{W}}}themeColor"),
                    }
                )
        result.update(
            {
                "tracked_nodes": tracked,
                "comment_markers": comments,
                "highlight_nodes": highlights,
                "rsid_attributes": rsids,
                "story_colors": colors,
                "non_grayscale_story_colors": [
                    color
                    for color in colors
                    if color["theme"]
                    or (
                        color["val"] not in (None, "auto")
                        and re.fullmatch(r"[0-9A-Fa-f]{6}", color["val"] or "")
                        and not (
                            color["val"][0:2].lower()
                            == color["val"][2:4].lower()
                            == color["val"][4:6].lower()
                        )
                    )
                ],
            }
        )

        document_root = ET.fromstring(archive.read("word/document.xml"))
        sections = []
        for sect in document_root.findall(".//w:sectPr", NS):
            pg = sect.find("w:pgSz", NS)
            margin = sect.find("w:pgMar", NS)
            cols = sect.find("w:cols", NS)
            sections.append(
                {
                    "page_w": pg.get(f"{{{W}}}w") if pg is not None else None,
                    "page_h": pg.get(f"{{{W}}}h") if pg is not None else None,
                    "orientation": pg.get(f"{{{W}}}orient") if pg is not None else None,
                    "margin_top": margin.get(f"{{{W}}}top") if margin is not None else None,
                    "margin_bottom": margin.get(f"{{{W}}}bottom") if margin is not None else None,
                    "margin_left": margin.get(f"{{{W}}}left") if margin is not None else None,
                    "margin_right": margin.get(f"{{{W}}}right") if margin is not None else None,
                    "columns": cols.get(f"{{{W}}}num") if cols is not None else "1",
                    "column_space": cols.get(f"{{{W}}}space") if cols is not None else None,
                }
            )
        result["sections"] = sections

        core_root = ET.fromstring(archive.read("docProps/core.xml"))
        result["metadata"] = {
            "creator": (core_root.findtext("dc:creator", default="", namespaces=NS) or ""),
            "last_modified_by": (core_root.findtext("cp:lastModifiedBy", default="", namespaces=NS) or ""),
            "title": (core_root.findtext("dc:title", default="", namespaces=NS) or ""),
            "subject": (core_root.findtext("dc:subject", default="", namespaces=NS) or ""),
            "description": (core_root.findtext("dc:description", default="", namespaces=NS) or ""),
        }
        if "docProps/app.xml" in names:
            app_root = ET.fromstring(archive.read("docProps/app.xml"))
            result["app_metadata"] = {
                "company": app_root.findtext("ep:Company", default="", namespaces=NS) or "",
                "manager": app_root.findtext("ep:Manager", default="", namespaces=NS) or "",
            }

    doc = Document(path)
    paragraph_text = [clean_cell(paragraph.text) for paragraph in doc.paragraphs]
    full_text = "\n".join(text for text in paragraph_text if text)
    result["paragraph_count"] = len(doc.paragraphs)
    result["table_count"] = len(doc.tables)
    result["table_shapes"] = [[len(table.rows), len(table.columns)] for table in doc.tables]
    result["text_characters"] = len(full_text)
    result["external_editing_instruction_hits"] = [
        phrase
        for phrase in ("蓝色代表", "已经改写", "复制改写", "Turnitin", "AI率", "祛AI")
        if phrase.lower() in (full_text + story_text).lower()
    ]
    result["identity_hits"] = {
        token: (full_text + "\n" + story_text).lower().count(token.lower())
        for token in (
            "Lin Zhanyi",
            "Hong Kong Metropolitan University",
            "lzzzy20041125@outlook.com",
            "Anonymous Authors",
        )
    }

    tables = [
        [[clean_cell(cell.text) for cell in row.cells] for row in table.rows]
        for table in doc.tables
    ]
    result["tables"] = tables
    expected_table_one = [
        ["Route", "Semantic content", "Candidate labels", "Endpoint"],
        ["Tree", "No", "No", "Tree"],
        ["Main", "Fixed T0-3B", "No", "Tree after accepted rescue or exact fallback"],
        ["NoSem", "Removed", "No", "Matched control tree"],
        ["Perm", "Row-permuted", "No", "Matched control tree"],
    ]
    expected_table_two = [
        ["Dataset", "Shot", "Seed", "Tree", "Main", "NoSem", "Perm", "M G/R", "N G/R", "P G/R"],
        *expected_episode_rows(data),
    ]
    expected_table_three = [
        ["Condition", "Threshold", "Observed", "Result"],
        *expected_decision_rows(data),
    ]
    expected_table_four = [
        ["Hypothesis", "Evidence needed to challenge the boundary"],
        [
            "Support gate is underpowered",
            "Pre-registered stronger support evidence activates stably on both datasets without label access",
        ],
        [
            "Semantic geometry misses tree errors",
            "A frozen alternative representation separates from both controls and improves the tree",
        ],
        [
            "Transfer is the bottleneck",
            "Accepted rescues are accurate and yield positive paired endpoint effects under the same anchor",
        ],
        [
            "Effect is distribution-specific",
            "A preselected broader grid passes dataset and shot consistency conditions",
        ],
    ]
    result["table_comparisons"] = {
        "table_i_exact": len(tables) >= 1 and tables[0] == expected_table_one,
        "table_ii_exact": len(tables) >= 2 and tables[1] == expected_table_two,
        "table_iii_exact": len(tables) >= 3 and tables[2] == expected_table_three,
        "table_iv_exact": len(tables) >= 4 and tables[3] == expected_table_four,
    }
    if len(tables) >= 2 and len(tables[1]) == 19:
        rows = tables[1][1:]
        result["ledger_derived"] = {
            "episode_rows": len(rows),
            "main_tree_equal_rows": sum(row[3] == row[4] for row in rows),
            "main_gate_total": sum(int(row[7].split("/")[0]) for row in rows),
            "main_rescue_total": sum(int(row[7].split("/")[1]) for row in rows),
            "nosem_gate_total": sum(int(row[8].split("/")[0]) for row in rows),
            "nosem_rescue_total": sum(int(row[8].split("/")[1]) for row in rows),
            "perm_gate_total": sum(int(row[9].split("/")[0]) for row in rows),
            "perm_rescue_total": sum(int(row[9].split("/")[1]) for row in rows),
        }
    if len(tables) >= 3:
        result["decision_derived"] = {
            "condition_rows": len(tables[2]) - 1,
            "pass_rows": sum(row[3] == "PASS" for row in tables[2][1:]),
            "fail_rows": sum(row[3] == "FAIL" for row in tables[2][1:]),
        }

    ref_start = next((idx for idx, text in enumerate(paragraph_text) if text == "REFERENCES"), None)
    before_refs = "\n".join(paragraph_text[:ref_start]) if ref_start is not None else full_text
    refs = paragraph_text[ref_start + 1 :] if ref_start is not None else []
    ref_numbers = [int(match.group(1)) for text in refs if (match := re.match(r"^\[(\d+)\]", text))]
    cited_numbers = sorted(expand_citations(before_refs))
    result["references"] = {
        "count": len(ref_numbers),
        "numbers": ref_numbers,
        "continuous_1_15": ref_numbers == list(range(1, 16)),
        "cited_numbers": cited_numbers,
        "every_reference_cited": set(range(1, 16)).issubset(cited_numbers),
        "unknown_citation_numbers": sorted(set(cited_numbers) - set(range(1, 16))),
    }
    result["fact_tokens"] = {
        token: full_text.count(token)
        for token in (
            "18 episodes",
            "72 method rows",
            "0 failure rows",
            "+0.0000000",
            "+0.0009668",
            "-0.0005258",
            "NO_GO",
        )
    }
    return result


def pdf_audit(path: Path) -> dict:
    reader = PdfReader(str(path))
    pages = []
    full_text = ""
    for page in reader.pages:
        text = page.extract_text() or ""
        full_text += "\n" + text
        box = page.mediabox
        pages.append(
            {
                "width": float(box.width),
                "height": float(box.height),
                "rotation": page.get("/Rotate", 0),
                "text_characters": len(text),
            }
        )
    return {
        "path": str(path),
        "sha256": sha256(path),
        "size": path.stat().st_size,
        "pages": len(reader.pages),
        "page_boxes": pages,
        "encrypted": reader.is_encrypted,
        "identity_hits": {
            token: full_text.lower().count(token.lower())
            for token in (
                "Lin Zhanyi",
                "Hong Kong Metropolitan University",
                "lzzzy20041125@outlook.com",
                "Anonymous Authors",
            )
        },
        "external_editing_instruction_hits": [
            phrase
            for phrase in ("蓝色代表", "已经改写", "复制改写", "Turnitin", "AI率", "祛AI")
            if phrase.lower() in full_text.lower()
        ],
        "fact_tokens": {
            token: full_text.count(token)
            for token in (
                "18 episodes",
                "72 method rows",
                "0 failure rows",
                "+0.0000000",
                "+0.0009668",
                "-0.0005258",
                "NO_GO",
            )
        },
    }


def main() -> None:
    data = validate_contract()
    pattern = "AdapTabPrompt_ICCEIC2026_Final_*"
    paths = sorted(ROOT.glob(pattern))
    report = {
        "frozen_report": data["report"],
        "files": [
            docx_audit(path, data) if path.suffix.lower() == ".docx" else pdf_audit(path)
            for path in paths
        ],
    }
    if "--summary" in sys.argv:
        compact_files = []
        for item in report["files"]:
            compact = {
                key: value
                for key, value in item.items()
                if key not in {"story_colors", "tables", "page_boxes"}
            }
            compact_files.append(compact)
        report = {"frozen_report": report["frozen_report"], "files": compact_files}
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
